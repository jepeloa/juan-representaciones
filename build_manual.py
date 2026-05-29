"""Build Manual de uso — Catálogo Juan.docx from screenshots.

Uses python-docx to assemble a polished manual with:
  - cover page
  - admin section
  - client section
  - mobile section
  - tech notes
"""
from __future__ import annotations
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent
SHOTS = ROOT / 'screenshots'
LOGO = ROOT / 'media' / 'logo_juan.png'
OUT = ROOT / 'Manual de uso — Catálogo Juan.docx'

# Paleta (deepened pastels)
COL_BRAND_700 = RGBColor(0x38, 0x4c, 0x60)
COL_BRAND_500 = RGBColor(0x55, 0x73, 0x90)
COL_SAGE_700 = RGBColor(0x48, 0x56, 0x4b)
COL_SAGE_500 = RGBColor(0x6e, 0x84, 0x78)
COL_BEIGE_600 = RGBColor(0x7a, 0x64, 0x40)
COL_MUTED = RGBColor(0x6f, 0x8c, 0xa6)
COL_TEXT = RGBColor(0x38, 0x4c, 0x60)
COL_BG = RGBColor(0xf7, 0xf3, 0xe8)


def set_cell_bg(cell, color_hex: str):
    """Set table cell background color."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_heading(doc: Document, text: str, level: int = 1, color: RGBColor = COL_BRAND_700):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.font.color.rgb = color
    run.font.name = 'Plus Jakarta Sans'
    if level == 0:
        run.font.size = Pt(32)
    elif level == 1:
        run.font.size = Pt(20)
    elif level == 2:
        run.font.size = Pt(15)
    else:
        run.font.size = Pt(12)
    run.font.bold = True
    return h


def add_paragraph(doc: Document, text: str, *, italic: bool = False,
                  color: RGBColor = COL_TEXT, size: int = 11, bold: bool = False, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = 'Inter'
    run.font.size = Pt(size)
    run.italic = italic
    run.bold = bold
    run.font.color.rgb = color
    return p


def add_callout(doc: Document, label: str, body: str, bg='efe8d3', accent_color=COL_BEIGE_600):
    """Add a colored callout box using a 1-cell table."""
    t = doc.add_table(rows=1, cols=1)
    t.autofit = True
    cell = t.rows[0].cells[0]
    set_cell_bg(cell, bg)
    p1 = cell.paragraphs[0]
    p1_run = p1.add_run(label.upper())
    p1_run.bold = True
    p1_run.font.size = Pt(9)
    p1_run.font.color.rgb = accent_color
    p1_run.font.name = 'Inter'
    p2 = cell.add_paragraph()
    p2_run = p2.add_run(body)
    p2_run.font.size = Pt(10.5)
    p2_run.font.color.rgb = COL_TEXT
    p2_run.font.name = 'Inter'


def add_screenshot(doc: Document, name: str, caption: str | None = None, width: float = 6.3):
    """Add a screenshot from SHOTS folder."""
    path = SHOTS / name
    if not path.exists():
        add_paragraph(doc, f'[Screenshot {name} no disponible]', italic=True, color=COL_MUTED, size=9)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap.add_run(caption)
        cap_run.italic = True
        cap_run.font.size = Pt(9.5)
        cap_run.font.color.rgb = COL_MUTED
        cap_run.font.name = 'Inter'


def add_step(doc: Document, num: int, title: str, body: str = ''):
    """Numbered step with colored number badge."""
    p = doc.add_paragraph()
    num_run = p.add_run(f'{num}. ')
    num_run.bold = True
    num_run.font.color.rgb = COL_BRAND_500
    num_run.font.size = Pt(11.5)
    num_run.font.name = 'Inter'
    title_run = p.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(11.5)
    title_run.font.color.rgb = COL_TEXT
    title_run.font.name = 'Inter'
    if body:
        body_p = doc.add_paragraph(style='List Continue')
        body_p.paragraph_format.left_indent = Inches(0.25)
        b_run = body_p.add_run(body)
        b_run.font.size = Pt(10.5)
        b_run.font.color.rgb = COL_TEXT
        b_run.font.name = 'Inter'


def page_break(doc: Document):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def main():
    doc = Document()

    # Adjust margins
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Inter'
    style.font.size = Pt(11)
    style.font.color.rgb = COL_TEXT

    # ===== COVER =====
    # Empty space top
    for _ in range(4):
        doc.add_paragraph()

    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO), width=Inches(2.2))

    add_paragraph(doc, '', size=8)
    add_heading(doc, 'Manual de uso', 0)
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = h.add_run('Catálogo Juan — gastronómico mayorista')
    sub.font.size = Pt(16)
    sub.font.color.rgb = COL_SAGE_700
    sub.italic = True
    sub.font.name = 'Inter'

    for _ in range(2):
        doc.add_paragraph()

    add_paragraph(doc, 'Versión 1.0', align=WD_ALIGN_PARAGRAPH.CENTER, color=COL_MUTED, size=10)
    add_paragraph(doc, '© 2026', align=WD_ALIGN_PARAGRAPH.CENTER, color=COL_MUTED, size=10)

    for _ in range(6):
        doc.add_paragraph()

    add_callout(
        doc,
        'Acceso',
        'http://24.199.118.218:8003/\n\n'
        'Administrador:  juan / juan2026\n'
        'Cliente demo:   cliente / cliente2026',
    )

    page_break(doc)

    # ===== TOC manual =====
    add_heading(doc, 'Contenido', 1)
    toc_items = [
        ('1', 'Introducción'),
        ('2', 'Acceso y roles'),
        ('3', 'Catálogo: explorar productos'),
        ('4', 'Carrito y generación de órdenes'),
        ('5', 'Administración'),
        ('5.1', 'Cargar y editar productos'),
        ('5.2', 'Condiciones de pago y términos'),
        ('5.3', 'Gestión de usuarios'),
        ('6', 'Uso desde el celular'),
        ('7', 'Notas técnicas'),
    ]
    for num, label in toc_items:
        p = doc.add_paragraph()
        r1 = p.add_run(f'{num}    ')
        r1.font.color.rgb = COL_BRAND_500
        r1.bold = True
        r1.font.name = 'Inter'
        r2 = p.add_run(label)
        r2.font.size = Pt(11)
        r2.font.color.rgb = COL_TEXT
        r2.font.name = 'Inter'

    page_break(doc)

    # ===== 1. INTRODUCCIÓN =====
    add_heading(doc, '1. Introducción', 1)
    add_paragraph(
        doc,
        'Catálogo Juan es una aplicación web para consultar el catálogo consolidado de proveedores '
        'mayoristas gastronómicos, armar órdenes de compra y enviarlas por email. Tiene dos roles: '
        'el administrador (que carga productos, gestiona condiciones de pago, usuarios y configura el '
        'envío de órdenes) y el cliente (que arma su orden y descarga el PDF).',
    )
    add_paragraph(doc, '')
    add_paragraph(doc, 'Características principales', bold=True, size=12, color=COL_BRAND_700)
    for feat in [
        '• Catálogo unificado con búsqueda, filtros y vista tabla / grilla.',
        '• Detalle de producto con galería de fotos.',
        '• Carrito con cálculo automático según condición de pago elegida.',
        '• Generación de orden en PDF y envío por email automático.',
        '• Panel de administración para cargar / editar productos, definir condiciones y usuarios.',
        '• Diseño responsive — funciona en escritorio, tablet y teléfono.',
    ]:
        add_paragraph(doc, feat, size=11)

    page_break(doc)

    # ===== 2. ACCESO Y ROLES =====
    add_heading(doc, '2. Acceso y roles', 1)
    add_paragraph(
        doc,
        'Ingresá a la aplicación con tu usuario y contraseña. La pantalla de login muestra los dos '
        'roles disponibles. El acceso es persistente: la sesión se mantiene abierta una semana.',
    )
    add_screenshot(doc, '01_login.png', 'Pantalla de login con el panel de marca y la tarjeta de roles.')

    add_heading(doc, 'Roles disponibles', 2, COL_SAGE_700)
    add_callout(
        doc,
        'Administrador (juan)',
        'Ve todas las páginas: catálogo, proveedores, cargar productos, condiciones de pago y gestión de usuarios. '
        'Puede crear, editar y eliminar productos, condiciones de pago, usuarios y configurar los textos del PDF '
        'y el email de notificación.',
        bg='e3ebf2',
        accent_color=COL_BRAND_500,
    )
    doc.add_paragraph()
    add_callout(
        doc,
        'Cliente (cliente)',
        'Ve catálogo, proveedores y "Mi orden". Puede armar pedidos, elegir condición de pago, agregar notas y '
        'generar la orden en PDF (que se envía automáticamente al email configurado por el administrador).',
        bg='eef2ee',
        accent_color=COL_SAGE_500,
    )

    page_break(doc)

    # ===== 3. CATÁLOGO =====
    add_heading(doc, '3. Catálogo: explorar productos', 1)
    add_paragraph(
        doc,
        'Después del login te llevamos directo al catálogo. Tenés dos vistas: tabla (compacta, ideal para '
        'comparar precios) y grilla (con fotos grandes).',
    )
    add_screenshot(doc, '02_admin_catalogo_table.png', 'Vista tabla: 660 productos en 12 proveedores.')

    add_heading(doc, 'Vista grilla', 2, COL_SAGE_700)
    add_paragraph(
        doc,
        'Hacé clic en el botón "Grilla" arriba a la derecha para cambiar a vista de tarjetas con foto grande.',
    )
    add_screenshot(doc, '03_admin_catalogo_grid.png', 'Vista grilla — ideal para explorar visualmente.')

    add_heading(doc, 'Filtros y búsqueda', 2, COL_SAGE_700)
    add_paragraph(
        doc,
        'La barra de filtros permite buscar por texto (nombre, código o descripción), proveedor, categoría, '
        'moneda y rango de precio. Los resultados se filtran en tiempo real.',
    )
    add_screenshot(doc, '04_admin_busqueda_fadeco.png', 'Búsqueda "fadeco" — 3 resultados con precios entre $901.700 y $1.097.000.')

    add_heading(doc, 'Detalle de producto', 2, COL_SAGE_700)
    add_paragraph(
        doc,
        'Hacé clic en cualquier fila del catálogo para abrir el detalle: galería de fotos, descripción completa, '
        'IVA, código de barras, cantidad por bulto y archivo de origen.',
    )
    add_screenshot(doc, '05_admin_product_detail_modal.png', 'Modal de detalle con galería de imágenes y todos los campos.')

    page_break(doc)

    # ===== 4. CARRITO Y ÓRDENES =====
    add_heading(doc, '4. Carrito y generación de órdenes', 1)
    add_paragraph(
        doc,
        'Para armar una orden de compra agregás productos al carrito, elegís la condición de pago y generás el PDF. '
        'El PDF se envía automáticamente al email configurado.',
    )

    add_heading(doc, 'Paso 1: Agregar productos', 2, COL_SAGE_700)
    add_paragraph(
        doc,
        'En el catálogo, cada producto tiene un botón "+" a la derecha. Al hacerle clic se agrega al carrito '
        'y el botón se transforma en un control de cantidad −/[cantidad]/+, para que puedas seguir sumando '
        'o restar unidades sin salir del catálogo.',
    )
    add_screenshot(doc, '06_admin_inline_quantity.png',
                   'Control inline de cantidad: se ven los productos ya agregados en color sage.')

    add_heading(doc, 'Paso 2: Ir a "Mi orden"', 2, COL_SAGE_700)
    add_paragraph(
        doc,
        'El ítem "Mi orden" en el menú lateral muestra un contador con la cantidad total de items. '
        'Hacé clic para ver el carrito.',
    )
    add_screenshot(doc, '22_cliente_carrito.png', 'Pantalla "Mi orden" con los productos seleccionados y la sidebar de condición de pago.')

    add_heading(doc, 'Paso 3: Ajustar cantidades', 2, COL_SAGE_700)
    add_paragraph(
        doc,
        'En el carrito podés sumar / restar unidades de cada producto con los botones +/− y quitarlo con el botón × '
        'arriba a la derecha de cada fila. El total se recalcula automáticamente.',
    )
    add_screenshot(doc, '23_cliente_carrito_qty.png', 'Cantidad actualizada y total recalculado en vivo.')

    add_heading(doc, 'Paso 4: Elegir condición de pago', 2, COL_SAGE_700)
    add_paragraph(
        doc,
        'Las condiciones de pago vienen configuradas por el administrador. Por defecto incluyen:',
    )
    for cond in [
        '• Contado — 5% de descuento sobre lista (× 0.95)',
        '• Cheque 30 días — precio de lista (× 1.00)',
        '• Cheque 60 días — lista + 8% (× 1.08)',
        '• Cheque 90 días — lista + 15% (× 1.15)',
    ]:
        add_paragraph(doc, cond, size=10.5)
    add_paragraph(doc, '')
    add_paragraph(
        doc,
        'Al cambiar la condición, los totales se recalculan multiplicando por el factor correspondiente.',
    )
    add_screenshot(doc, '24_cliente_carrito_condicion.png',
                   'Selector de condición de pago y notas opcionales para el pedido.')

    add_heading(doc, 'Paso 5: Generar orden y PDF', 2, COL_SAGE_700)
    add_paragraph(
        doc,
        'Hacé clic en "Crear orden y generar PDF". La orden se guarda en el sistema, se envía un email al '
        'administrador con el PDF adjunto y la app te muestra un confirmador con el estado del envío.',
    )
    add_screenshot(doc, '25_cliente_orden_creada.png',
                   'Orden creada. Botones para abrir o descargar el PDF, y estado del envío automático por email.')

    add_callout(
        doc,
        'Sobre el email automático',
        'Cuando el cliente genera la orden, el sistema envía un email con el PDF al destinatario configurado '
        'en /admin/condiciones. El cliente ve el estado del envío en pantalla: enviado, pendiente, '
        'no configurado o falló. Si falla puede reenviar más tarde.',
    )

    page_break(doc)

    # ===== 5. ADMINISTRACIÓN =====
    add_heading(doc, '5. Administración', 1)
    add_paragraph(
        doc,
        'Las secciones de administración solo son visibles para usuarios con rol administrador. Aparecen '
        'en el menú lateral como "Cargar productos", "Condiciones" y "Usuarios".',
    )

    # 5.1 PRODUCTOS
    add_heading(doc, '5.1 Cargar y editar productos', 2, COL_SAGE_700)
    add_paragraph(
        doc,
        'Desde "Cargar productos" tenés acceso a la tabla completa con todos los productos del catálogo. '
        'Podés buscar por texto, filtrar por proveedor, y editar o eliminar cualquier producto (incluyendo '
        'los importados automáticamente desde Excel).',
    )
    add_screenshot(doc, '07_admin_productos_lista.png', 'Listado de productos con búsqueda y filtro por proveedor.')

    add_heading(doc, 'Crear un producto nuevo', 3)
    add_step(doc, 1, 'Hacé clic en "Nuevo producto" (arriba a la derecha).')
    add_step(doc, 2, 'Elegí proveedor del desplegable, o usá "+ Nuevo" para crear uno al instante.')
    add_step(doc, 3, 'Opcional: elegí o creá una categoría dentro del proveedor.')
    add_step(doc, 4, 'Completá nombre, código (SKU), descripción, precio, moneda, IVA, bulto y código de barras.')
    add_step(doc, 5, 'Subí una o varias fotos: se redimensionan automáticamente a 600×600 px.')
    add_step(doc, 6, 'Guardá. El producto queda disponible en el catálogo inmediatamente.')
    add_screenshot(doc, '08_admin_productos_form_nuevo.png', 'Formulario vacío para crear un producto nuevo.')
    add_screenshot(doc, '09_admin_productos_form_lleno.png', 'Formulario con todos los datos del producto.')
    add_screenshot(doc, '10_admin_productos_creado.png', 'Producto recién creado en la lista.')

    add_heading(doc, 'Editar un producto existente', 3)
    add_paragraph(
        doc,
        'Buscá el producto en la lista, hacé clic en "Editar", modificá los campos que necesites y guardá. '
        'Si querés agregar más fotos, subí nuevas; si querés quitar la foto actual marcá el checkbox '
        '"Quitar la foto actual al guardar".',
    )
    add_screenshot(doc, '11_admin_productos_editar.png', 'Modo edición con la foto actual visible y opción de quitarla.')
    add_screenshot(doc, '12_admin_productos_actualizado.png', 'Producto actualizado con el nuevo precio.')

    page_break(doc)

    # 5.2 CONDICIONES
    add_heading(doc, '5.2 Condiciones de pago y términos', 2, COL_SAGE_700)
    add_paragraph(
        doc,
        'En "Condiciones" definís cómo se calculan los precios según el plazo de pago y configurás los textos '
        'que aparecen en el PDF de cada orden.',
    )
    add_screenshot(doc, '13_admin_condiciones.png', 'Listado de condiciones de pago activas.')

    add_heading(doc, 'Agregar una condición de pago', 3)
    add_step(doc, 1, 'Hacé clic en "Nueva condición".')
    add_step(doc, 2, 'Nombre (ej. "Cheque 45 días").')
    add_step(doc, 3, 'Descripción opcional (ej. "Lista + 5%").')
    add_step(doc, 4, 'Multiplicador: 1.00 = precio lista. 1.10 = +10%. 0.95 = −5%.')
    add_step(doc, 5, 'Días (opcional) y orden de aparición.')
    add_step(doc, 6, 'Guardá. Aparece inmediatamente disponible para los clientes en su carrito.')
    add_screenshot(doc, '14_admin_condiciones_form.png', 'Formulario de condición con badge automático del porcentaje.')

    add_heading(doc, 'Términos del PDF y email destinatario', 3)
    add_paragraph(
        doc,
        'Abajo del listado de condiciones hay un panel con los textos del PDF y el email destinatario.',
    )
    for item in [
        '• Nombre de la empresa: aparece en el header del PDF.',
        '• Contacto: teléfono o email que aparece debajo del nombre.',
        '• Aclaración: texto que sale al pie del PDF (ej. "Precios estimativos, sujetos a confirmación").',
        '• Términos: condiciones generales del catálogo.',
        '• Email para recibir las órdenes nuevas: dirección a la que se envía el PDF cuando un cliente crea una orden.',
    ]:
        add_paragraph(doc, item, size=10.5)
    add_callout(
        doc,
        'Email de notificación',
        'Cuando el cliente genera una orden, se envía automáticamente al destinatario configurado. Si el campo '
        'queda vacío, las órdenes se guardan pero no se mandan emails (el cliente igual puede descargar el PDF).',
        bg='e3ebf2',
        accent_color=COL_BRAND_500,
    )

    page_break(doc)

    # 5.3 USUARIOS
    add_heading(doc, '5.3 Gestión de usuarios', 2, COL_SAGE_700)
    add_paragraph(
        doc,
        'En "Usuarios" creás, editás y eliminás cuentas. Cada usuario puede ser administrador o cliente.',
    )
    add_screenshot(doc, '16_admin_usuarios.png', 'Tabla de usuarios con rol y estado (activo / inactivo).')
    add_heading(doc, 'Operaciones disponibles', 3)
    for op in [
        '• Crear: botón "Nuevo usuario" — pedirá usuario, nombre completo, contraseña y rol.',
        '• Editar: cambiá nombre, contraseña, rol y estado.',
        '• Activar / desactivar: hacé clic en el badge "Activo / Inactivo".',
        '• Eliminar: con confirmación. No podés eliminar tu propia cuenta.',
    ]:
        add_paragraph(doc, op, size=10.5)

    # PROVEEDORES
    add_heading(doc, 'Proveedores', 2, COL_SAGE_700)
    add_paragraph(
        doc,
        'En "Proveedores" ves la lista de los 12 proveedores integrados, con conteo de productos por cada uno. '
        'Hacé clic en cualquiera para filtrar el catálogo por ese proveedor.',
    )
    add_screenshot(doc, '19_proveedores.png', 'Grilla de proveedores con conteo de productos.')

    page_break(doc)

    # ===== 6. MOBILE =====
    add_heading(doc, '6. Uso desde el celular', 1)
    add_paragraph(
        doc,
        'La aplicación es 100% responsive y funciona en cualquier celular o tablet. En pantallas chicas (< 768 px) '
        'el menú lateral se convierte en un drawer deslizante y las tablas se transforman en cards stackeadas, '
        'para que sea cómodo el uso con el dedo.',
    )

    add_heading(doc, 'Login y catálogo', 2, COL_SAGE_700)
    # Layout en grid 2 columnas para mobile
    t = doc.add_table(rows=1, cols=2)
    t.autofit = True
    c1 = t.rows[0].cells[0]
    c2 = t.rows[0].cells[1]
    c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    c1.paragraphs[0].add_run().add_picture(str(SHOTS / '30_mobile_login.png'), width=Inches(2.6))
    c2.paragraphs[0].add_run().add_picture(str(SHOTS / '31_mobile_catalogo.png'), width=Inches(2.6))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_r = cap.add_run('Login y catálogo en mobile (cards en lugar de tabla).')
    cap_r.italic = True
    cap_r.font.size = Pt(9.5)
    cap_r.font.color.rgb = COL_MUTED

    add_heading(doc, 'Menú lateral (drawer)', 2, COL_SAGE_700)
    add_paragraph(
        doc,
        'El ícono de hamburguesa arriba a la izquierda abre el menú lateral. Tocá afuera o el botón × para cerrarlo. '
        'La barra superior también tiene un acceso directo al carrito con badge del contador.',
    )
    add_screenshot(doc, '32_mobile_sidebar.png', 'Menú lateral abierto en mobile con backdrop.', width=3.5)

    add_heading(doc, 'Agregar productos al carrito', 2, COL_SAGE_700)
    add_paragraph(
        doc,
        'Cada producto en la lista mobile tiene un botón "+" de 44 px (tap-friendly). Tras agregar, se convierte '
        'en el control −/cantidad/+ para que puedas sumar más unidades sin volver al carrito.',
    )
    add_screenshot(doc, '33_mobile_catalogo_carrito.png',
                   'Producto agregado: el botón "+" se transformó en control de cantidad.',
                   width=3.5)

    add_heading(doc, 'Carrito y orden', 2, COL_SAGE_700)
    t = doc.add_table(rows=1, cols=2)
    t.autofit = True
    c1 = t.rows[0].cells[0]
    c2 = t.rows[0].cells[1]
    c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    c1.paragraphs[0].add_run().add_picture(str(SHOTS / '34_mobile_carrito.png'), width=Inches(2.6))
    c2.paragraphs[0].add_run().add_picture(str(SHOTS / '35_mobile_orden_creada.png'), width=Inches(2.6))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_r = cap.add_run('Carrito mobile y confirmación tras crear la orden.')
    cap_r.italic = True
    cap_r.font.size = Pt(9.5)
    cap_r.font.color.rgb = COL_MUTED

    page_break(doc)

    # ===== 7. NOTAS TÉCNICAS =====
    add_heading(doc, '7. Notas técnicas', 1)

    add_heading(doc, 'Arquitectura', 2, COL_SAGE_700)
    for item in [
        '• Frontend: Angular 18 (standalone components) + Tailwind CSS, servido por nginx.',
        '• Backend: FastAPI + SQLAlchemy 2 + Alembic, sobre MySQL 8.',
        '• PDF: generado con ReportLab (logo, tabla por moneda, totales, disclaimer).',
        '• Email: SMTP estándar vía smtplib (Python stdlib), background task de FastAPI.',
        '• Despliegue: Docker Compose con 3 servicios (mysql, backend, frontend).',
    ]:
        add_paragraph(doc, item, size=10.5)

    add_heading(doc, 'Configurar el envío de emails', 2, COL_SAGE_700)
    add_paragraph(
        doc,
        'El destinatario se configura desde la web en /admin/condiciones. Las credenciales SMTP se cargan '
        'como variables de entorno en el archivo /srv/catalogo-app/.env del servidor:',
    )
    p = doc.add_paragraph()
    code_run = p.add_run(
        'SMTP_HOST=smtp.gmail.com\n'
        'SMTP_PORT=587\n'
        'SMTP_USER=tu_cuenta@gmail.com\n'
        'SMTP_PASSWORD=app_password_de_gmail\n'
        'SMTP_FROM=tu_cuenta@gmail.com\n'
        'SMTP_USE_TLS=true'
    )
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(9.5)
    code_run.font.color.rgb = COL_TEXT
    add_paragraph(
        doc,
        'Para Gmail hay que generar un App Password (no la clave normal). Después correr '
        '"docker compose up -d" en el servidor para que el backend tome los valores.',
        size=10.5,
    )

    add_heading(doc, 'Reiniciar / actualizar', 2, COL_SAGE_700)
    add_paragraph(doc, 'Conectado por SSH al servidor:', size=10.5)
    p = doc.add_paragraph()
    code_run = p.add_run(
        'ssh root@24.199.118.218\n'
        'cd /srv/catalogo-app\n'
        'docker compose ps                       # ver estado\n'
        'docker compose restart backend          # reiniciar backend\n'
        'docker compose logs -f backend          # ver logs en vivo'
    )
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(9.5)

    add_heading(doc, 'Backup', 2, COL_SAGE_700)
    add_paragraph(
        doc,
        'La base de datos vive en el volumen Docker "catalogo-app_mysql_data". Las fotos subidas por admin '
        'se guardan en /srv/catalogo-data/uploads/ del host. Para backup completo: dump del volumen MySQL '
        '+ rsync del directorio /srv/catalogo-data/.',
        size=10.5,
    )

    # ===== END =====
    page_break(doc)
    for _ in range(8):
        doc.add_paragraph()
    add_paragraph(
        doc,
        '— Fin del manual —',
        align=WD_ALIGN_PARAGRAPH.CENTER,
        color=COL_MUTED,
        italic=True,
        size=11,
    )

    doc.save(str(OUT))
    print(f'\n✓ Manual generado: {OUT}')
    print(f'  Tamaño: {OUT.stat().st_size / 1024:.1f} KB')


if __name__ == '__main__':
    main()
